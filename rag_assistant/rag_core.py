from __future__ import annotations

import json
import math
import os
import re
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree

import numpy as np
import requests
import torch
from bs4 import BeautifulSoup
from docx import Document
from transformers import AutoModel, AutoTokenizer


PROJECT_DIR = Path(__file__).resolve().parent
COURSE_ROOT = PROJECT_DIR.parent
COURSE_MODEL_PATH = COURSE_ROOT / "open_models" / "bge-m3"
LOCAL_MODEL_PATH = PROJECT_DIR / "models" / "bge-m3"
DEFAULT_MODEL_PATH = str(COURSE_MODEL_PATH if COURSE_MODEL_PATH.exists() else LOCAL_MODEL_PATH)


@dataclass
class Chunk:
    """一段可检索文本，以及它来自哪个文件。"""

    id: str
    source: str
    title: str
    text: str


def read_text_file(path: Path) -> str:
    """读取普通文本文件，尽量兼容 UTF-8 和中文 Windows 编码。"""

    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def read_docx(path: Path) -> str:
    """读取 Word 文档中的段落和表格文字。"""

    doc = Document(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_md(path: Path) -> str:
    """用标准库读取 Markdown 里的文本。

    Markdown 本质上是一个 zip 包，演示材料文字放在 md/lessons/slide*.xml 中。
    这里不用 python-md，方便在课堂机器上少装依赖。
    """

    slide_texts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            name
            for name in zf.namelist()
            if name.startswith("md/lessons/slide") and name.endswith(".xml")
        )

        for index, name in enumerate(slide_names, start=1):
            xml = zf.read(name)
            root = ElementTree.fromstring(xml)
            texts = [
                node.text.strip()
                for node in root.iter()
                if node.tag.endswith("}t") and node.text and node.text.strip()
            ]
            if texts:
                slide_texts.append(f"【第 {index} 页】\n" + "\n".join(texts))

    return "\n\n".join(slide_texts)


def clean_text(text: str) -> str:
    """做一点轻量清洗，保留中文标点和教程里的结构。"""

    text = BeautifulSoup(text, "html.parser").get_text("\n")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def load_documents(source_dir: Path) -> list[tuple[Path, str]]:
    """从课程目录读取可作为知识库的文件。

    默认跳过离线镜像、临时缓存和代码依赖说明，避免把无关内容放进知识库。
    """

    supported = {".md", ".txt", ".docx", ".md"}
    skip_parts = {
        "offline_bundle",
        "__pycache__",
        ".venv",
        "storage",
        "node_modules",
        "open_models",
    }
    documents: list[tuple[Path, str]] = []

    for path in sorted(source_dir.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in supported:
            continue
        if any(part in skip_parts for part in path.parts):
            continue
        if path.name.lower() in {"requirements.txt"}:
            continue

        try:
            if path.suffix.lower() in {".md", ".txt"}:
                text = read_text_file(path)
            elif path.suffix.lower() == ".docx":
                text = read_docx(path)
            elif path.suffix.lower() == ".md":
                text = read_md(path)
            else:
                continue
        except Exception as exc:
            print(f"[跳过] {path}: {exc}")
            continue

        text = clean_text(text)
        if len(text) >= 80:
            documents.append((path, text))

    return documents


def split_text(text: str, chunk_size: int = 650, overlap: int = 120) -> list[str]:
    """把长文档切成带 overlap 的小块。

    overlap 的作用是保留相邻片段之间的上下文，避免答案刚好被切断。
    """

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 2 <= chunk_size:
            current = f"{current}\n\n{paragraph}".strip()
            continue

        if current:
            chunks.append(current)

        if len(paragraph) <= chunk_size:
            current = paragraph
        else:
            start = 0
            while start < len(paragraph):
                end = start + chunk_size
                chunks.append(paragraph[start:end])
                start = max(end - overlap, start + 1)
            current = ""

    if current:
        chunks.append(current)

    if overlap <= 0 or len(chunks) <= 1:
        return chunks

    merged: list[str] = []
    for index, chunk in enumerate(chunks):
        if index == 0:
            merged.append(chunk)
            continue
        prev_tail = chunks[index - 1][-overlap:]
        merged.append(f"{prev_tail}\n{chunk}")
    return merged


def build_chunks(source_dir: Path, chunk_size: int = 650, overlap: int = 120) -> list[Chunk]:
    """读取目录并切片，生成 RAG 的基础知识单元。"""

    chunks: list[Chunk] = []
    for path, text in load_documents(source_dir):
        relative = path.relative_to(source_dir)
        title = path.stem
        for position, chunk_text in enumerate(split_text(text, chunk_size, overlap), start=1):
            chunks.append(
                Chunk(
                    id=f"{len(chunks):06d}",
                    source=str(relative),
                    title=title,
                    text=chunk_text,
                )
            )
    return chunks


class BGEEmbedder:
    """用本地 BGE-m3 生成文本向量。"""

    def __init__(self, model_path: str = DEFAULT_MODEL_PATH, device: str = "cpu"):
        self.model_path = model_path
        self.device = device
        self.tokenizer = AutoTokenizer.from_pretrained(model_path)
        self.model = AutoModel.from_pretrained(model_path)
        self.model.to(device)
        self.model.eval()

    @torch.no_grad()
    def encode(self, texts: list[str], batch_size: int = 4) -> np.ndarray:
        """把多段文本转成归一化向量。"""

        vectors: list[np.ndarray] = []
        for start in range(0, len(texts), batch_size):
            batch = texts[start : start + batch_size]
            encoded = self.tokenizer(
                batch,
                padding=True,
                truncation=True,
                max_length=512,
                return_tensors="pt",
            )
            encoded = {key: value.to(self.device) for key, value in encoded.items()}
            output = self.model(**encoded)
            last_hidden = output.last_hidden_state
            mask = encoded["attention_mask"].unsqueeze(-1).float()
            pooled = (last_hidden * mask).sum(dim=1) / mask.sum(dim=1).clamp(min=1e-6)
            pooled = torch.nn.functional.normalize(pooled, p=2, dim=1)
            vectors.append(pooled.cpu().numpy())

        return np.vstack(vectors).astype("float32")


def tokenize_for_keyword(text: str) -> list[str]:
    """非常轻量的关键词切分，适合中英文混合课堂材料。"""

    text = text.lower()
    words = re.findall(r"[a-zA-Z0-9_]{2,}|[\u4e00-\u9fff]{2,}", text)
    chinese_chars = re.findall(r"[\u4e00-\u9fff]", text)
    bigrams = ["".join(chinese_chars[i : i + 2]) for i in range(max(0, len(chinese_chars) - 1))]
    return words + bigrams


def keyword_score(query: str, text: str) -> float:
    """关键词检索分数：查询词命中越多，分数越高。"""

    query_terms = tokenize_for_keyword(query)
    if not query_terms:
        return 0.0
    text_terms = set(tokenize_for_keyword(text))
    hits = sum(1 for term in query_terms if term in text_terms)
    return hits / math.sqrt(len(query_terms) + 1)


class LocalRAGIndex:
    """本地 RAG 索引：保存 chunks 和向量，负责检索。"""

    def __init__(self, chunks: list[Chunk], embeddings: np.ndarray, embedder: BGEEmbedder):
        self.chunks = chunks
        self.embeddings = embeddings
        self.embedder = embedder

    @classmethod
    def load(cls, storage_dir: Path, model_path: str = DEFAULT_MODEL_PATH) -> "LocalRAGIndex":
        chunks_data = json.loads((storage_dir / "chunks.json").read_text(encoding="utf-8"))
        chunks = [Chunk(**item) for item in chunks_data]
        embeddings = np.load(storage_dir / "embeddings.npy")
        embedder = BGEEmbedder(model_path=model_path)
        return cls(chunks, embeddings, embedder)

    def search(self, query: str, top_k: int = 5, vector_weight: float = 0.75) -> list[dict]:
        """混合检索：向量相似度 + 关键词命中。

        vector_weight 越高，越相信语义相似；越低，越相信关键词精确命中。
        """

        query_vector = self.embedder.encode([query], batch_size=1)[0]
        vector_scores = self.embeddings @ query_vector
        keyword_scores = np.array(
            [keyword_score(query, chunk.text) for chunk in self.chunks],
            dtype="float32",
        )

        if keyword_scores.max() > 0:
            keyword_scores = keyword_scores / keyword_scores.max()

        final_scores = vector_weight * vector_scores + (1 - vector_weight) * keyword_scores
        top_indices = np.argsort(final_scores)[::-1][:top_k]

        results: list[dict] = []
        for rank, index in enumerate(top_indices, start=1):
            chunk = self.chunks[int(index)]
            results.append(
                {
                    "rank": rank,
                    "score": float(final_scores[index]),
                    "vector_score": float(vector_scores[index]),
                    "keyword_score": float(keyword_scores[index]),
                    "source": chunk.source,
                    "title": chunk.title,
                    "text": chunk.text,
                }
            )
        return results


def save_index(chunks: list[Chunk], embeddings: np.ndarray, storage_dir: Path) -> None:
    storage_dir.mkdir(parents=True, exist_ok=True)
    (storage_dir / "chunks.json").write_text(
        json.dumps([chunk.__dict__ for chunk in chunks], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    np.save(storage_dir / "embeddings.npy", embeddings)


def build_and_save_index(
    source_dir: Path,
    storage_dir: Path,
    model_path: str = DEFAULT_MODEL_PATH,
    chunk_size: int = 650,
    overlap: int = 120,
) -> None:
    start = time.time()
    chunks = build_chunks(source_dir, chunk_size=chunk_size, overlap=overlap)
    if not chunks:
        raise RuntimeError(f"没有从 {source_dir} 读取到可用文档。")

    print(f"读取并切分完成：{len(chunks)} 个片段")
    embedder = BGEEmbedder(model_path=model_path)
    embeddings = embedder.encode([chunk.text for chunk in chunks], batch_size=4)
    save_index(chunks, embeddings, storage_dir)
    print(f"索引已保存到：{storage_dir}")
    print(f"总耗时：{time.time() - start:.1f} 秒")


def build_context(search_results: list[dict], max_chars: int = 4500) -> str:
    """把检索结果整理成喂给大模型的上下文。"""

    parts: list[str] = []
    total = 0
    for item in search_results:
        text = item["text"].strip()
        block = (
            f"[资料 {item['rank']}] 来源：{item['source']}\n"
            f"相关分数：{item['score']:.3f}\n"
            f"{text}"
        )
        if total + len(block) > max_chars:
            break
        parts.append(block)
        total += len(block)
    return "\n\n---\n\n".join(parts)


def ask_deepseek(
    question: str,
    context: str,
    api_key: str | None = None,
    model: str = "deepseek-chat",
    base_url: str = "https://api.deepseek.com",
    timeout: int = 60,
) -> str:
    """调用 DeepSeek，要求模型只基于检索资料回答。"""

    api_key = api_key or os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise RuntimeError("未设置 DEEPSEEK_API_KEY。请先在环境变量中配置 DeepSeek API Key。")

    url = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    system_message = (
        "你是面向培训课堂的 RAG 问答助手。"
        "回答要自然、清楚、偏实操，不要编造资料中没有的信息。"
        "如果资料不足，请直接说明缺少依据，并给出下一步建议。"
    )
    user_message = f"""请基于下面的课程资料回答问题。

【课程资料】
{context}

【问题】
{question}

回答要求：
1. 先给出直接答案。
2. 如果涉及步骤，请用简洁的分点说明。
3. 末尾列出你参考了哪些资料来源。
"""

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_message},
        ],
        "temperature": 0.2,
        "stream": False,
    }

    response = requests.post(url, headers=headers, json=payload, timeout=timeout)
    response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"].strip()